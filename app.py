import io
import os
import uuid
import base64
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

# Import python-docx untuk generate Word berlogo
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from ai_engine import (
    generate_quiz_batch, get_ai_hint_stream, get_ai_solution_stream,
    create_table_if_not_exists, update_progress_siswa, init_db_connection,
    generate_lkpd_content
)

st.set_page_config(
    page_title="RoboMANTAP-AI (Assistant Intelligence)",
    page_icon="logo.png",
    layout="wide",
    initial_sidebar_state="auto"
)

# Inisialisasi Tabel Database saat aplikasi pertama kali dimuat
create_table_if_not_exists()

components.html(
    """
    <script>
    if ('wakeLock' in navigator) {
        let wakeLock = null;
        const requestWakeLock = async () => {
            try {
                wakeLock = await navigator.wakeLock.request('screen');
            } catch (err) {
                console.log(`${err.name}, ${err.message}`);
            }
        };
        requestWakeLock();
        document.addEventListener('visibilitychange', async () => {
            if (wakeLock !== null && document.visibilityState === 'visible') {
                await requestWakeLock();
            }
        });
    }
    </script>
    """,
    height=0,
)

# Custom Styling (Dari Kode Awal - Dipertahankan 100%)
st.markdown("""
    <style>
    .mode-card {
        background-color: var(--secondary-background-color);
        color: var(--text-color);
        border: 1px solid rgba(128, 128, 128, 0.2);
        padding: 24px;
        border-radius: 12px;
        text-align: center;
        margin-bottom: 15px;
    }
    .mapel-card {
        background-color: var(--secondary-background-color);
        border: 1px solid rgba(5, 150, 105, 0.3);
        padding: 16px;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 10px;
    }
    .school-header {
        background: linear-gradient(135deg, #064e3b 0%, #022c22 100%);
        border: 1px solid #059669;
        border-radius: 12px;
        padding: 16px;
        text-align: center;
        margin-bottom: 20px;
        box-shadow: 0 4px 20px rgba(5, 150, 105, 0.15);
    }
    .school-title { color: #ffffff; font-weight: 800; font-size: 13px; margin: 0; }
    .school-subtitle { color: #6ee7b7; font-size: 11px; margin-top: 4px; font-weight: 500; }
    .stButton>button { width: 100%; min-height: 48px; font-size: 16px !important; border-radius: 8px !important; }
    
    /* Paksa Warna Tombol Utama Menjadi Hijau Emerald MANTAP */
    div.stButton > button[kind="primary"],
    div.stButton > button {
        background-color: #059669 !important;
        background-image: none !important;
        color: #ffffff !important;
        border: 1px solid #047857 !important;
        width: 100%;
        min-height: 48px;
        font-size: 16px !important;
        font-weight: 600 !important;
        border-radius: 8px !important;
        transition: all 0.3s ease;
    }
    
    /* Efek Hover Tombol saat Diarahkan Kursor */
    div.stButton > button[kind="primary"]:hover,
    div.stButton > button:hover {
        background-color: #047857 !important;
        border-color: #065f46 !important;
        box-shadow: 0 4px 12px rgba(5, 150, 105, 0.4) !important;
    }

    .guru-card {
        background: linear-gradient(135deg, #1e3a8a 0%, #172554 100%);
        color: white;
        border: 1px solid #3b82f6;
        padding: 18px;
        border-radius: 12px;
        text-align: center;
        margin-bottom: 15px;
    }
    </style>
""", unsafe_allow_html=True)

# Helper Base64 Image
def get_image_base64(path):
    if os.path.exists(path):
        with open(path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    return None

logo_mantap_b64 = get_image_base64("logo.png")
logo_nexus_b64 = get_image_base64("nexus_logo.png")

img_mantap_html = f'<img src="data:image/png;base64,{logo_mantap_b64}" style="height: 70px; margin-bottom: 8px;">' if logo_mantap_b64 else '<div style="font-size: 32px;">🎓</div>'

# Session States Initialization
if "page" not in st.session_state: st.session_state.page = "landing"
if "jenjang" not in st.session_state: st.session_state.jenjang = None
if "mapel" not in st.session_state: st.session_state.mapel = None
if "stage" not in st.session_state: st.session_state.stage = "Internal"
if "selected_submateri" not in st.session_state: st.session_state.selected_submateri = []
if "quiz_data" not in st.session_state: st.session_state.quiz_data = []
if "user_answers" not in st.session_state: st.session_state.user_answers = {}
if "current_index" not in st.session_state: st.session_state.current_index = 0

# Penambahan State untuk Integrasi Database Siswa & Guru
if "nama_siswa" not in st.session_state: st.session_state.nama_siswa = ""
if "session_id" not in st.session_state: st.session_state.session_id = str(uuid.uuid4())
if "guru_auth" not in st.session_state: st.session_state.guru_auth = False
if "ai_hint_cache" not in st.session_state: st.session_state.ai_hint_cache = {}
if "ai_solution_cache" not in st.session_state: st.session_state.ai_solution_cache = {}

# Database Kisi-Kisi Operasional OMI 2026
KISI_KISI_OMI = {
    "MTs (Sederajat SMP)": {
        "Matematika": ["Bilangan", "Aljabar", "Aritmetika Sosial", "Geometri", "Peluang", "Statistika", "Perbandingan & Proporsi", "Problem Solving", "Konteks OMI (Keislaman & Sains)"],
        "IPA Terintegrasi": ["Makhluk Hidup & Sel", "Sistem Organ", "Genetika & Keanekaragaman", "Ekologi", "Zat & Perubahannya", "Energi & Kalor", "Gerak & Gaya", "Getaran, Gelombang & Optik", "Listrik & Kemagnetan", "Bumi & Antariksa", "Eksperimen & Data", "Konteks OMI"],
        "IPS Terintegrasi": ["Geografi", "Kependudukan", "Ekonomi", "Sejarah Indonesia", "Sejarah Islam", "Sosial & Budaya", "Kewarganegaraan", "Lingkungan & Pembangunan", "Literasi Data", "Konteks OMI"]
    },
    "MA (Sederajat SMA)": {
        "Matematika Terintegrasi": ["Bilangan & Teori Bilangan", "Aljabar & Fungsi", "Geometri", "Kombinatorika & Peluang", "Statistika", "Problem Solving", "Konteks OMI"],
        "Biologi Terintegrasi": ["Sel & Biokimia", "Genetika", "Fisiologi", "Botani & Zoologi", "Ekologi", "Evolusi & Keanekaragaman", "Bioteknologi & Lingkungan", "Konteks OMI"],
        "Fisika Terintegrasi": ["Mekanika", "Fluida", "Getaran & Gelombang", "Optik", "Suhu & Kalor", "Listrik & Magnet", "Fisika Modern", "Eksperimen & Data", "Konteks OMI"],
        "Kimia Terintegrasi": ["Struktur Atom & Periodik", "Ikatan Kimia", "Stoikiometri", "Larutan & Asam-Basa", "Redoks", "Termokimia & Kinetika", "Kesetimbangan", "Organik & Lingkungan", "Konteks OMI"],
        "Ekonomi Terintegrasi": ["Ekonomi Dasar", "Mikroekonomi", "Makroekonomi", "Kebijakan Ekonomi", "Akuntansi", "Pasar Modal & Keuangan", "Ekonomi Digital", "Ekonomi Islam", "Analisis Data"],
        "Geografi Terintegrasi": ["Peta & Keruangan", "Geologi & Geomorfologi", "Atmosfer & Iklim", "Hidrosfer", "Biosfer", "Kependudukan", "Sumber Daya & Lingkungan", "Bencana", "SIG & Data Spasial", "Konteks OMI"]
    }
}

# Header Utama
st.markdown(f"""
<div class="school-header">
    <div style="text-align: center;">
        {img_mantap_html}
    </div>
    <div class="school-title">MA DAN MTs AL IRSYAD AL ISLAMIYYAH BONDOWOSO</div>
    <div class="school-subtitle">
        Madrasah Aliyah dan Tsanawiyah Al Irsyad Putri Bondowoso (MANTAP) &nbsp;•&nbsp; 
        <span style="color: #6ee7b7; font-weight: 600;">Powered by RoboMANTAP-AI (Assistant Intelligence)</span>
    </div>
</div>
""", unsafe_allow_html=True)

# Sidebar Control
with st.sidebar:
    st.markdown("""
    <div style="background: linear-gradient(135deg, #064e3b 0%, #022c22 100%); padding: 16px; border-radius: 12px; border: 1px solid #059669; text-align: center; margin-bottom: 15px; box-shadow: 0 4px 12px rgba(0,0,0,0.15);">
        <div style="font-size: 26px; margin-bottom: 4px;">🧕🏼</div>
        <div style="color: #ffffff; font-weight: 700; font-size: 16px; letter-spacing: 0.5px;">RoboMANTAP-AI</div>
        <div style="color: #6ee7b7; font-size: 11px; font-weight: 500; margin-bottom: 6px;">Assistant Intelligence System</div>
        <div style="font-size: 10px; color: #a7f3d0; opacity: 0.85; border-top: 1px solid rgba(255,255,255,0.15); padding-top: 4px; font-style: italic;">Engineered by U.Project Nexus</div>
    </div>
    """, unsafe_allow_html=True)

    if st.session_state.page not in ["guru_login", "guru_dashboard"]:
        st.markdown("""
        <div style="background: var(--secondary-background-color); border: 1px solid rgba(5, 150, 105, 0.3); padding: 12px 14px; border-radius: 10px; margin-bottom: 15px;">
            <div style="display: flex; align-items: center; justify-content: space-between; margin-bottom: 6px;">
                <span style="font-size: 11px; font-weight: 600; opacity: 0.7;">ENGINE STATUS</span>
                <span style="font-size: 10px; background: #059669; color: white; padding: 2px 8px; border-radius: 12px; font-weight: 700;">LIVE 🟢</span>
            </div>
            <div style="font-size: 11px; line-height: 1.6; opacity: 0.9;">
                ⚡ <b>Model:</b> U.Project Nexus Intelligence v3.6<br>
                🎯 <b>Core:</b> Bina Prestasi OMI 2026<br>
                ⏱️ <b>Response:</b> Real-time AI
            </div>
        </div>
        """, unsafe_allow_html=True)

        if st.session_state.jenjang and st.session_state.mapel:
            st.markdown(f"""
            <div style="background: rgba(5, 150, 105, 0.08); border-left: 4px solid #059669; padding: 10px 12px; border-radius: 6px; margin-bottom: 15px;">
                <div style="font-size: 10px; opacity: 0.6; text-transform: uppercase; font-weight: 700;">Sesi Aktif</div>
                <div style="font-size: 12px; font-weight: 700; color: var(--text-color);">{st.session_state.mapel}</div>
                <div style="font-size: 11px; opacity: 0.8;">{st.session_state.jenjang} • {st.session_state.stage}</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("""
        <div style="background: var(--secondary-background-color); border: 1px solid rgba(128,128,128,0.2); padding: 12px 14px; border-radius: 10px; margin-bottom: 15px;">
            <div style="font-size: 11px; font-weight: 700; opacity: 0.8; margin-bottom: 8px;">📋 ATURAN SKORING CBT</div>
            <div style="display: flex; justify-content: space-between; font-size: 11px; margin-bottom: 4px;">
                <span>✅ Jawaban Benar</span>
                <b style="color: #059669;">+4 Poin</b>
            </div>
            <div style="display: flex; justify-content: space-between; font-size: 11px; margin-bottom: 4px;">
                <span>❌ Jawaban Salah</span>
                <b style="color: #ef4444;">-1 Poin</b>
            </div>
            <div style="display: flex; justify-content: space-between; font-size: 11px;">
                <span>⚪ Tidak Dijawab</span>
                <b style="opacity: 0.6;">0 Poin</b>
            </div>
        </div>
        """, unsafe_allow_html=True)

    if st.button("🏠 Kembali ke Beranda Utama", use_container_width=True):
        st.session_state.page = "landing"
        st.session_state.jenjang = None
        st.session_state.mapel = None
        st.session_state.guru_auth = False
        st.rerun()

    sidebar_nexus_html = f'<div style="background: #ffffff; padding: 6px 14px; border-radius: 10px; box-shadow: 0 4px 10px rgba(0,0,0,0.08); display: inline-block; margin-bottom: 8px; border: 1px solid rgba(0,0,0,0.05);"><img src="data:image/png;base64,{logo_nexus_b64}" style="height: 42px; max-width: 100%; display: block; margin: 0 auto;"></div>' if logo_nexus_b64 else ''
    
    st.markdown(f"""
    <div style="text-align: center; margin-top: 20px; padding-top: 15px; border-top: 1px dashed rgba(128,128,128,0.2);">
        <div style="font-size: 10px; opacity: 0.7; margin-bottom: 6px; text-transform: uppercase; letter-spacing: 0.5px;">Engineered by</div>
        {sidebar_nexus_html}
        <div style="font-size: 11px; opacity: 0.85; line-height: 1.3;">
            <b style="color: var(--text-color);">U.Project Nexus System</b><br>
            <span style="font-size: 10px; opacity: 0.7;">AI Integration & B2B Solutions</span><br>
            <span style="font-size: 9px; opacity: 0.5;">&copy; 2026 All Rights Reserved</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

#helper LKPD
def set_cell_background(cell, hex_color):
    """Memberikan warna latar belakang pada sel tabel Word."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_lkpd_docx_buffer(mapel, kelas, topik, ai_content, logo_path="logo.png"):
    """
    Membuat file Word (.docx) Interaktif, Berwarna, Elegan khas Al-Irsyad Bondowoso.
    """
    doc = Document()

    # Set Margin Halaman
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # --------------------------------------------------------------------------
    # 1. KOP SURAT BERLOGO (ELEGAN)
    # --------------------------------------------------------------------------
    table_kop = doc.add_table(rows=1, cols=2)
    table_kop.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_kop.columns[0].width = Inches(1.1)
    table_kop.columns[1].width = Inches(5.4)

    # Logo Sekolah
    if os.path.exists(logo_path):
        p_logo = table_kop.cell(0, 0).paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(logo_path, width=Inches(0.95))

    # Teks Kop
    p_text = table_kop.cell(0, 1).paragraphs[0]
    p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r1 = p_text.add_run("LEMBAR KERJA PESERTA DIDIK (LKPD) INTERAKTIF\n")
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(5, 150, 105) # Emerald Green

    r2 = p_text.add_run("Madrasah Aliyah dan Tsanawiyah Al-Irsyad Al-Islamiyah Putri Bondowoso\n")
    r2.bold = True
    r2.font.size = Pt(10)
    
    r3 = p_text.add_run("Model Pembelajaran HOTS & Integrasi Nilai Keislaman")
    r3.italic = True
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("―" * 58)

    # --------------------------------------------------------------------------
    # 2. KARD IDENTITAS MATERI & SANTRI (CARD SHADED)
    # --------------------------------------------------------------------------
    table_meta = doc.add_table(rows=1, cols=1)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_meta = table_meta.cell(0, 0)
    set_cell_background(cell_meta, "ECFDF5") # Soft Mint Green Fill

    p_meta = cell_meta.paragraphs[0]
    p_meta.paragraph_format.space_before = Pt(6)
    p_meta.paragraph_format.space_after = Pt(6)
    
    r_meta_title = p_meta.add_run("📌 IDENTITAS PEMBELAJARAN\n")
    r_meta_title.bold = True
    r_meta_title.font.size = Pt(10)
    r_meta_title.font.color.rgb = RGBColor(4, 120, 87)

    p_meta.add_run(f"• Mata Pelajaran\t: {mapel}\n").bold = True
    p_meta.add_run(f"• Kelas / Jenjang\t: {kelas}\n").bold = True
    p_meta.add_run(f"• Topik Utama\t: {topik}\n").bold = True
    p_meta.add_run(f"• Nama Santri/Kelompok : ...........................................................")

    doc.add_paragraph() # Spacing

    # --------------------------------------------------------------------------
    # 3. BAGIAN A: TUJUAN PEMBELAJARAN
    # --------------------------------------------------------------------------
    t_a = doc.add_table(rows=1, cols=1)
    cell_a = t_a.cell(0, 0)
    set_cell_background(cell_a, "059669") # Emerald Banner
    p_a = cell_a.paragraphs[0]
    r_a = p_a.add_run("🎯 [A] TUJUAN PEMBELAJARAN")
    r_a.bold = True
    r_a.font.color.rgb = RGBColor(255, 255, 255)

    tujuan_list = ai_content.get("tujuan", [])
    for idx, t in enumerate(tujuan_list, 1):
        p_t = doc.add_paragraph()
        p_t.add_run(f"{idx}. {t}").font.size = Pt(9.5)

    # --------------------------------------------------------------------------
    # 4. BAGIAN B: APERSEPSI & KONSEP (KOTAK BIRU SOFT)
    # --------------------------------------------------------------------------
    t_b_head = doc.add_table(rows=1, cols=1)
    set_cell_background(t_b_head.cell(0, 0), "059669")
    r_b_head = t_b_head.cell(0, 0).paragraphs[0].add_run("📖 [B] APERSEPSI & EXPLORASI KONSEP")
    r_b_head.bold = True
    r_b_head.font.color.rgb = RGBColor(255, 255, 255)

    t_b_box = doc.add_table(rows=1, cols=1)
    cell_b_box = t_b_box.cell(0, 0)
    set_cell_background(cell_b_box, "F0F9FF") # Light Sky Blue Fill
    p_b_content = cell_b_box.paragraphs[0]
    p_b_content.add_run(ai_content.get("ringkasan", "")).font.size = Pt(9.5)

    doc.add_paragraph()

    # --------------------------------------------------------------------------
    # 5. BAGIAN C: TUGAS EKSPLORASI & KOTAK JAWABAN SANTRI
    # --------------------------------------------------------------------------
    t_c_head = doc.add_table(rows=1, cols=1)
    set_cell_background(t_c_head.cell(0, 0), "059669")
    r_c_head = t_c_head.cell(0, 0).paragraphs[0].add_run("✍️ [C] TUGAS EKSPLORASI MANDIRI (HOTS)")
    r_c_head.bold = True
    r_c_head.font.color.rgb = RGBColor(255, 255, 255)

    # Soal 1 + Kotak Lembar Jawab
    p_s1 = doc.add_paragraph()
    p_s1.add_run(f"Soal 1: {ai_content.get('soal_1', '')}").bold = True
    
    t_ans1 = doc.add_table(rows=1, cols=1)
    cell_ans1 = t_ans1.cell(0, 0)
    set_cell_background(cell_ans1, "F9FAFB") # Soft Gray Answer Box
    p_ans1 = cell_ans1.paragraphs[0]
    p_ans1.add_run("Lembar Jawaban / Analisis Santri:\n\n\n\n").italic = True
    p_ans1.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # Soal 2 + Kotak Lembar Jawab
    p_s2 = doc.add_paragraph()
    p_s2.add_run(f"Soal 2: {ai_content.get('soal_2', '')}").bold = True

    t_ans2 = doc.add_table(rows=1, cols=1)
    cell_ans2 = t_ans2.cell(0, 0)
    set_cell_background(cell_ans2, "F9FAFB")
    p_ans2 = cell_ans2.paragraphs[0]
    p_ans2.add_run("Lembar Jawaban / Analisis Santri:\n\n\n\n").italic = True
    p_ans2.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph()

    # --------------------------------------------------------------------------
    # 6. BAGIAN D: REFLEKSI KEISLAMAN (KOTAK EMAS SOFT)
    # --------------------------------------------------------------------------
    t_d_head = doc.add_table(rows=1, cols=1)
    set_cell_background(t_d_head.cell(0, 0), "D97706") # Amber/Gold Banner
    r_d_head = t_d_head.cell(0, 0).paragraphs[0].add_run("🌿 [D] REFLEKSI KEISLAMAN & HIKMAH SANTRI")
    r_d_head.bold = True
    r_d_head.font.color.rgb = RGBColor(255, 255, 255)

    t_d_box = doc.add_table(rows=1, cols=1)
    cell_d_box = t_d_box.cell(0, 0)
    set_cell_background(cell_d_box, "FEF3C7") # Warm Yellow Fill
    p_d_content = cell_d_box.paragraphs[0]
    r_d_text = p_d_content.add_run(f'"{ai_content.get("refleksi", "")}"')
    r_d_text.italic = True
    r_d_text.font.size = Pt(9.5)
    r_d_text.font.color.rgb = RGBColor(146, 64, 14)

    # Save to Buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
# ==============================================================================
# 1. TAMPILAN AWAL (GERBANG SISWA & GURU)
# ==============================================================================
if st.session_state.page == "landing":
    st.markdown("<h2 style='font-size: 25px; text-align: center;'>🏆 BINA PRESTASI OMI 2026</h2>", unsafe_allow_html=True)
    st.markdown("<p style='font-size: 12px; text-align: center; opacity: 0.8;'>Pilih Jenjang Pendidikan untuk Memulai Pembinaan Olimpiade</p>", unsafe_allow_html=True)
    st.write("---")

    st.markdown("#### 📝 Mulai Latihan CBT")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
        <div class="mode-card">
            <h2>🏫 TINGKAT MTs</h2>
            <p style="opacity: 0.7; font-size: 14px;">Madrasah Tsanawiyah Al-Irsyad Putri</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Masuk Modul MTs ➔", key="btn_mts", use_container_width=True, type="primary"):
            st.session_state.jenjang = "MTs (Sederajat SMP)"
            st.session_state.page = "select_mapel"
            st.rerun()

    with col2:
        st.markdown("""
        <div class="mode-card">
            <h2>🏛️ TINGKAT MA</h2>
            <p style="opacity: 0.7; font-size: 14px;">Madrasah Aliyah Al-Irsyad Putri</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Masuk Modul MA ➔", key="btn_ma", use_container_width=True, type="primary"):
            st.session_state.jenjang = "MA (Sederajat SMA)"
            st.session_state.page = "select_mapel"
            st.rerun()

    st.write("---")
    st.markdown("#### 🧕🏼 Portal GuruMANTAP")
    st.markdown("""
    <div class="guru-card">
        <h2 style="margin:0; font-size: 20px;">🔴 Live Monitoring & AI Generator</h2>
        <p style="font-size: 10px; opacity:0.8; margin-top:5px;">Pantau skor siswa secara real-time, generate soal, dan integrasi WhatsApp</p>
    </div>
    """, unsafe_allow_html=True)
    if st.button("🔒 Masuk Portal Guru ➔", use_container_width=True):
        st.session_state.page = "guru_login"
        st.rerun()

# ==============================================================================
# 2. LOGIN GURU & DASHBOARD
# ==============================================================================
elif st.session_state.page == "guru_login":
    st.subheader("🔒 Akses Portal GuruMANTAP")

    st.markdown("""
    <div style="background-color: rgba(28, 131, 225, 0.1); border-left: 4px solid #1c83e1; padding: 10px 12px; border-radius: 6px; font-size: 10px; color: var(--text-color); margin-bottom: 15px;">
        ⚠️ Fitur ini dilindungi PIN untuk menjaga kerahasiaan nilai siswa dan soal CBT
    </div>
    """, unsafe_allow_html=True)
    
    pin_input = st.text_input("Masukkan PIN Akses:", type="password", placeholder="**********")
    if st.button("Login", type="primary"):
        if pin_input == "MANTAP2026":
            st.session_state.guru_auth = True
            st.session_state.page = "guru_dashboard"
            st.rerun()
        else:
            st.error("PIN Salah. Silakan coba lagi.")

elif st.session_state.page == "guru_dashboard":
    if not st.session_state.guru_auth:
        st.warning("Akses Ditolak.")
        st.stop()

    st.markdown("<p style='font-size: 27px; font-weight: bold; margin-bottom: 8px;'>🖥️ Dashboard GuruMANTAP</p>", unsafe_allow_html=True)
    tab1, tab2, tab3 = st.tabs(["🔴 Live Monitoring", "🧕Bank Soal", "📲 WA Automation"])

    with tab1:
        # Fungsi HTML untuk visual bar ala Quizizz (Disesuaikan jadi 10 Soal)
        def render_progress_bar_html(detail_list):
            if not isinstance(detail_list, list): return ""
            html = '<div style="display: flex; gap: 4px; align-items: center;">'
            for i in range(10): # Total soal per kuis = 10
                if i < len(detail_list):
                    val = detail_list[i]
                    color = "#10b981" if val is True else ("#ef4444" if val is False else "#d1d5db")
                else:
                    color = "#f3f4f6"
                html += f'<div style="background-color:{color}; height:14px; flex:1; border-radius:3px;"></div>'
            html += '</div>'
            return html

        # Auto-refresh setiap 3 detik
        @st.fragment(run_every="3s")
        def render_live_monitoring():
            conn = init_db_connection()
            if not conn:
                st.warning("Menunggu koneksi Database terhubung untuk Live Monitoring...")
                return

            try:
                df = conn.query("SELECT nama_siswa, jenjang, mapel, soal_sekarang, detail_jawaban, nilai_akhir, status FROM sesi_ujian ORDER BY updated_at DESC", ttl=0)
                
                if df.empty:
                    st.info("Belum ada siswa yang sedang mengerjakan ujian saat ini.")
                    return

                # KPI Metrics
                c1, c2, c3 = st.columns(3)
                c1.metric("Siswa Aktif", len(df[df['status'] == 'BERJALAN']))
                c2.metric("Sesi Selesai", len(df[df['status'] == 'SELESAI']))
                c3.metric("Rata-Rata Nilai", f"{df['nilai_akhir'].mean():.1f}")

                st.write("---")
                st.markdown("#### 🟢 Live Tracking Pengerjaan")
                
                # Render UI Tabel Kustom
                for index, row in df.iterrows():
                    with st.container():
                        col_nama, col_mapel, col_skor, col_bar = st.columns([3, 2, 1, 4])
                        status_badge = "✅" if row['status'] == 'SELESAI' else "🔄"
                        col_nama.markdown(f"**{row['nama_siswa']}** {status_badge}")
                        col_mapel.caption(f"{row['mapel']} ({row['jenjang'][:3]})")
                        col_skor.markdown(f"**Skor: {row['nilai_akhir']}**")
                        
                        # Render Kotak Hijau/Merah
                        try:
                            detail_list = row['detail_jawaban']
                            if isinstance(detail_list, str):
                                import json
                                detail_list = json.loads(detail_list)
                            col_bar.markdown(render_progress_bar_html(detail_list), unsafe_allow_html=True)
                        except:
                            col_bar.write("-")
                        st.divider()

            except Exception as e:
                st.error(f"Gagal mengambil data dari database: {e}")

        render_live_monitoring()

    with tab2:
        # Header Ringkas HP
        st.markdown("<p style='font-size: 18px; font-weight: bold; margin-bottom: 6px;'>🧕🏼 AI Quiz Generator & Analisis</p>", unsafe_allow_html=True)
        
        # Kotak Info Kustom Ringkas (12px)
        st.markdown("""
        <div style="background-color: #eff6ff; border-left: 4px solid #3b82f6; padding: 10px 12px; border-radius: 8px; font-size: 12px; line-height: 1.5; color: #1e3a8a; margin-bottom: 15px;">
            <b style="font-size: 13px;">🚀 Fitur Mendatang:</b>
            <ul style="margin: 6px 0 0 0; padding-left: 18px;">
                <li><b>Server Eksklusif:</b> Menghadirkan Server System terbaru U.Project Nexus Intelligence v3.6</li>
                <li><b>Generator Massal:</b> Buat puluhan/ratusan paket soal HOTS & tematik secara instan</li>
                <li><b>Export Cetak & PDF:</b> Siap cetak ber-template eksklusif sekolah</li>
                <li><b>Analisis Butir Soal:</b> Evaluasi daya pembeda & tingkat kesukaran soal</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)

    with tab3:
        # Header Ringkas HP
        st.markdown("<p style='font-size: 18px; font-weight: bold; margin-bottom: 6px;'>📲 WhatsApp Integration Engine</p>", unsafe_allow_html=True)
        
        # Kotak Info Kustom Ringkas (12px)
        st.markdown("""
        <div style="background-color: #eff6ff; border-left: 4px solid #3b82f6; padding: 10px 12px; border-radius: 8px; font-size: 12px; line-height: 1.5; color: #1e3a8a; margin-bottom: 15px;">
            <b style="font-size: 13px;">⚡ One-Click Automation By U.Project Nexus:</b>
            <ul style="margin: 6px 0 0 0; padding-left: 18px;">
                <li><b>Broadcast Hasil Ujian:</b> Laporan nilai otomatis ke WA orang tua & siswa</li>
                <li><b>ChatBot RoboMANTAP 24/7:</b> Asisten tutor pribadi siswa di rumah</li>
                <li><b>Auto-LKPD Guru:</b> Buat LKPD otomatis sesuai template khas sekolah</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)

        st.divider()

        # ==============================================================================
        # GENERATOR LKPD WORD (.DOCX) EKSKLUSIF BERLOGO
        # ==============================================================================
        st.markdown("""
        <div style="background: linear-gradient(135deg, #064e3b 0%, #022c22 100%); padding: 12px; border-radius: 10px; border: 1px solid #059669; color: white; margin-bottom: 12px;">
            <div style="font-size: 14px; font-weight: 700;">📄 Generator LKPD Word Berlogo Resmi</div>
            <div style="font-size: 11px; opacity: 0.85; margin-top: 2px;">Terintegrasi Gemini AI & Format Madrasah Al-Irsyad Al-Islamiyah Putri Bondowoso</div>
        </div>
        """, unsafe_allow_html=True)

        topic_lkpd = st.text_input("Topik / Materi Pembelajaran:", placeholder="Contoh: Persamaan Kuadrat / Tajwid Hukum Nun Mati", key="lkpd_topic")
        
        col_lkpd1, col_lkpd2 = st.columns(2)
        with col_lkpd1:
            kelas_lkpd = st.selectbox("Kelas / Jenjang:", ["VII MTs", "VIII MTs", "IX MTs", "X MA", "XI MA", "XII MA"], key="lkpd_kelas")
        with col_lkpd2:
            mapel_lkpd = st.selectbox("Mata Pelajaran:", ["Matematika", "IPA Terintegrasi", "IPS Terintegrasi", "PAI & Bahasa Arab", "Fisika", "Biologi", "Kimia"], key="lkpd_mapel")

        if st.button("📄 Generate LKPD Word (.docx) Berlogo", type="primary", use_container_width=True):
            if not topic_lkpd.strip():
                st.warning("⚠️ Ketik topik/materi pembelajarannya dulu ya!")
            else:
                with st.spinner("U.Project Nexus sedang menyusun dokumen Word berlogo resmi..."):
                    ai_content = generate_lkpd_content(mapel_lkpd, kelas_lkpd, topic_lkpd)
                    
                    if not ai_content:
                        st.error("Gagal menyusun LKPD. Silakan klik tombol sekali lagi.")
                    else:
                        # Buat File Word berlogo di Memory
                        docx_buffer = create_lkpd_docx_buffer(mapel_lkpd, kelas_lkpd, topic_lkpd, ai_content)
                        
                        st.success("✅ Dokumen LKPD Word (.docx) Berlogo Berhasil Diproduksi!")
                        
                        # Tombol Unduh Dokumen Word (.docx)
                        st.download_button(
                            label="📥 Download Dokumen LKPD (.docx)",
                            data=docx_buffer,
                            file_name=f"LKPD_AlIrsyad_{mapel_lkpd}_{topic_lkpd.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )



# ==============================================================================
# 3. TAMPILAN PILIHAN MATA PELAJARAN OMI 2026 (SISWA)
# ==============================================================================
elif st.session_state.page == "select_mapel":
    st.markdown(f"### 📚 Pilih Bidang OMI 2026 -<br><span style='color: #059669; display: inline-block;'>{st.session_state.jenjang}</span>", unsafe_allow_html=True)
    if st.button("⬅️ Kembali Pilih Jenjang"):
        st.session_state.page = "landing"
        st.rerun()

    st.write("---")
    mapel_dict = KISI_KISI_OMI[st.session_state.jenjang]
    cols = st.columns(len(mapel_dict))

    for idx, (mapel_name, submateri_list) in enumerate(mapel_dict.items()):
        with cols[idx % len(cols)]:
            st.markdown(f"""
            <div class="mapel-card">
                <h4>{mapel_name}</h4>
                <p style="font-size: 12px; opacity: 0.7;">{len(submateri_list)} Submateri Operasional</p>
            </div>
            """, unsafe_allow_html=True)
            if st.button(f"Pilih {mapel_name}", key=f"btn_mapel_{idx}", type="primary", use_container_width=True):
                st.session_state.mapel = mapel_name
                st.session_state.page = "setup"
                st.rerun()

# ==============================================================================
# 4. SETUP CBT & BIODATA SISWA
# ==============================================================================
elif st.session_state.page == "setup":
    st.markdown(f"""
    <div style="font-size: 23px; font-weight: bold; line-height: 1.4; margin-bottom: 10px;">
        ⚙️ Persiapan CBT:<br>
        <span style="font-size: 17px; color: #059669; font-weight: 600;">
            {st.session_state.mapel} ({st.session_state.jenjang})
        </span>
    </div>
    """, unsafe_allow_html=True)
    if st.button("⬅️ Ganti Mata Pelajaran"):
        st.session_state.page = "select_mapel"
        st.rerun()

    st.write("---")
    # Input Biodata Siswa Wajib
    st.markdown("#### 📝 Masukkan Data Diri Kamu")
    st.session_state.nama_siswa = st.text_input("Nama Lengkap:", value=st.session_state.nama_siswa, placeholder="Contoh: Fulanah binti Fulan")
    st.write("---")

    c1, c2 = st.columns([5, 7])

    with c1:
        st.subheader("1. Konfigurasi Ujian")
        st.session_state.stage = st.radio("Pilih Tahap Pembinaan:", ["Internal", "Kab/Kota", "Provinsi", "Nasional"])
        
        available_submateri = KISI_KISI_OMI[st.session_state.jenjang][st.session_state.mapel]
        st.session_state.selected_submateri = st.multiselect(
            "Pilih Submateri: (Click (Select all) untuk memilih semua Submateri)",
            available_submateri,
            default=[],
            placeholder="Pilih submateri di sini..."
        )

    with c2:
        st.subheader("2. Petunjuk CBT RoboMANTAP")
        st.markdown("""
        * **Jumlah Soal:** TEPAT 10 Soal Pilihan Ganda Terintegrasi per Sesi.
        * **Standar Pembinaan:** Mengacu Juknis OMI 2026 (Sains, Keislaman, & Literasi Data).
        * **Skoring:** Benar (+4), Salah (-1), Kosong (0).
        """)
        st.write("")
        if st.button("🚀 MARI MULAI SESI TEST SEKARANG!", type="primary", use_container_width=True):
            if not st.session_state.nama_siswa.strip():
                st.error("⚠️ Isi nama lengkap kamu dulu ya sebelum mulai!")
            else:
                st.session_state.session_id = str(uuid.uuid4())
                with st.spinner(f"RoboMANTAP sedang merancang 10 soal {st.session_state.mapel} Kamu. Tunggu sebentar ya... (nggak lama kok, hanya butuh waktu sekitar 15 detik saja! 😊)"):
                    quiz = generate_quiz_batch(
                        st.session_state.jenjang,
                        st.session_state.mapel,
                        st.session_state.stage,
                        st.session_state.selected_submateri
                    )
                    if quiz and len(quiz) == 10:
                        st.session_state.quiz_data = quiz
                        st.session_state.user_answers = {}
                        st.session_state.current_index = 0
                        # Sinkronisasi awal ke DB (Semua Kosong)
                        update_progress_siswa(
                            st.session_state.session_id, st.session_state.nama_siswa,
                            st.session_state.jenjang, st.session_state.mapel, 1, [], "BERJALAN"
                        )
                        st.session_state.page = "quiz"
                        st.rerun()
                    else:
                        st.error("Gagal membuat paket soal. Silakan klik tombol sekali lagi.")

# ==============================================================================
# 5. ENGINE TEST INTERAKTIF (10 SOAL CBT) + LIVE SYNC DATABASE
# ==============================================================================
elif st.session_state.page == "quiz":
    quiz_data = st.session_state.quiz_data
    curr_idx = st.session_state.current_index
    q = quiz_data[curr_idx]

    col_h1, col_h2 = st.columns([8, 4])
    with col_h1:
        st.subheader(f"📝 CBT OMI: {st.session_state.mapel} ({st.session_state.stage})")
        st.caption(f"👤 Siswa: **{st.session_state.nama_siswa}**")
    with col_h2:
        st.progress((curr_idx + 1) / 10)
        st.caption(f"Soal **{curr_idx + 1}** dari **10**")

    st.write("---")
    st.markdown(f"#### **Soal No. {curr_idx + 1}**")
    st.markdown(q["question"])
    st.write("")

    opts = q["options"]
    saved_ans = st.session_state.user_answers.get(curr_idx, None)
    default_opt_idx = opts.index(saved_ans) if saved_ans in opts else None

    selected_option = st.radio("Pilih Jawaban Anda:", opts, index=default_opt_idx, key=f"radio_q_{curr_idx}")
    
    # Trigger sinkronisasi jika ada pilihan jawaban yang berubah
    if selected_option and selected_option != saved_ans:
        st.session_state.user_answers[curr_idx] = selected_option
        
        # Hitung Array Status Jawaban (True/False/None) untuk DB (10 soal)
        detail = []
        for i in range(10):
            u_ans = st.session_state.user_answers.get(i, None)
            if u_ans is None:
                detail.append(None)
            else:
                is_correct = (u_ans == quiz_data[i]["correct_answer"])
                detail.append(is_correct)
                
        # Update Real-Time ke DB
        update_progress_siswa(
            st.session_state.session_id, st.session_state.nama_siswa,
            st.session_state.jenjang, st.session_state.mapel, curr_idx + 1, detail, "BERJALAN"
        )

    st.write("---")
    
    col_nav1, col_nav2, col_nav3 = st.columns([3, 6, 3])
    # Swap posisi Next/Submit ke atas untuk aksesibilitas HP
    with col_nav1:
        if curr_idx < 9:
            if st.button("Berikutnya ➡️", type="primary", use_container_width=True):
                st.session_state.current_index += 1
                st.rerun()
        else:
            if st.button("🏁 SUBMIT & SELESAIKAN", type="primary", use_container_width=True):
                st.session_state.page = "result"
                st.rerun()
    with col_nav3:
        if curr_idx > 0:
            if st.button("⬅️ Sebelumnya", use_container_width=True):
                st.session_state.current_index -= 1
                st.rerun()

    # Expander Hint dengan Live Streaming Text
    with st.expander("Kamu bingung? Konsultasi di sini sama aku, RoboMANTAP! 🧕🏼"):
        st.caption("Fungsi kolom ini: Tulis ide awal atau rumus yang mau kamu coba, nanti RoboMANTAP bakal kasih petunjuk jalan keluarnya tanpa langsung bocorin jawaban!")
        
        attempt_input = st.text_input(
            "Gagasan / Ide Logika Kamu apa coba?:",
            placeholder="Contoh: Aku masih belum bisa mengartikan Bahasa Arab. 😟",
            key=f"hint_in_{curr_idx}"
        )
        
        # Cache hint berdasarkan soal + ide siswa + mata pelajaran.
        hint_key = (
            st.session_state.mapel,
            curr_idx,
            q["question"],
            attempt_input.strip(),
        )

        if attempt_input.strip() and hint_key in st.session_state.ai_hint_cache:
            st.markdown("🧕🏼 **RoboMANTAP:**")
            st.markdown(st.session_state.ai_hint_cache[hint_key])

        if st.button("Diskusikan Yuk!", key=f"btn_hint_{curr_idx}"):
            if attempt_input.strip():
                if hint_key in st.session_state.ai_hint_cache:
                    st.markdown("🧕🏼 **RoboMANTAP:**")
                    st.markdown(st.session_state.ai_hint_cache[hint_key])
                else:
                    st.markdown("🧕🏼 **RoboMANTAP:**")
                    streamed_hint = st.write_stream(
                        get_ai_hint_stream(
                            q["question"],
                            attempt_input,
                            st.session_state.mapel,
                        ),
                        cursor="▌",
                    )

                    # Simpan hanya respons yang berhasil selesai tanpa pesan error.
                    if streamed_hint and "⚠️" not in str(streamed_hint):
                        st.session_state.ai_hint_cache[hint_key] = str(streamed_hint)
            else:
                st.info("💡 Tolong ketik sedikit ide kamu dulu ya, biar RoboMANTAP bisa kasih petunjuk yang pas!")


# ==============================================================================
# 6. SCORECARD & EVALUASI SESI
# ==============================================================================
elif st.session_state.page == "result":
    st.subheader(f"📊 Evaluasi CBT: {st.session_state.mapel} ({st.session_state.jenjang})")
    quiz_data = st.session_state.quiz_data
    user_answers = st.session_state.user_answers

    benar, salah, kosong, total_skor = 0, 0, 0, 0
    detail = []
    for idx, q in enumerate(quiz_data):
        u_ans = user_answers.get(idx, None)
        if u_ans is None:
            kosong += 1
            detail.append(None)
        elif u_ans == q["correct_answer"]:
            benar += 1; total_skor += 4
            detail.append(True)
        else:
            salah += 1; total_skor -= 1
            detail.append(False)

    # Sinkronisasi Final Status SELESAI ke DB Guru
    update_progress_siswa(
        st.session_state.session_id, st.session_state.nama_siswa,
        st.session_state.jenjang, st.session_state.mapel, 10, detail, "SELESAI"
    )

    if total_skor >= 32:
        feedback_msg = f"🌟 **Luar Biasa! (Skor: {total_skor}/40)**\n\nRoboMANTAP bangga banget sama kamu! Pemahaman kamu di materi {st.session_state.mapel} sudah sangat tajam. Pertahankan fokus kamu untuk Persiapan OMI 2026 ya! 🚀✨"
        feedback_type = "success"
    elif total_skor >= 16:
        feedback_msg = f"👍 **Kerja Bagus! (Skor: {total_skor}/40)**\n\nUsaha yang mantap! Kamu sudah paham sebagian besar konsepnya. Coba cek pembahasan di bawah untuk memperbaiki sedikit kekeliruan tadi ya! 💪😊"
        feedback_type = "info"
    else:
        feedback_msg = f"🌱 **Tetap Semangat! (Skor: {total_skor}/40)**\n\nJangan berkecil hati ya! Setiap kesalahan adalah proses belajar. Yuk pelajari pembahasan rinci di bawah dan coba latihan 10 soal lagi bersama RoboMANTAP! 🧕🏼❤️"
        feedback_type = "warning"

    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Total Skor CBT", f"{total_skor} / 40")
    k2.metric("Benar (+4)", f"{benar}")
    k3.metric("Salah (-1)", f"{salah}")
    k4.metric("Kosong (0)", f"{kosong}")

    if feedback_type == "success":
        st.success(f"🧕🏼 **Pesan dari RoboMANTAP:**\n\n{feedback_msg}")
    elif feedback_type == "info":
        st.info(f"🧕🏼 **Pesan dari RoboMANTAP:**\n\n{feedback_msg}")
    else:
        st.warning(f"🧕🏼 **Pesan dari RoboMANTAP:**\n\n{feedback_msg}")

    st.write("---")
    col_act1, col_act2 = st.columns(2)
    with col_act1:
        if st.button("🔄 LATIHAN SOAL LAGI DONG! (SESI BARU)", type="primary", use_container_width=True):
            st.cache_data.clear()
            with st.spinner("Sabar ya, RoboMANTAP sedang menyiapkan soal baru Kamu.. (nggak lama kok, hanya butuh waktu sekitar 15 detik saja! 😊)"):
                new_quiz = generate_quiz_batch(st.session_state.jenjang, st.session_state.mapel, st.session_state.stage, st.session_state.selected_submateri)
                if new_quiz and len(new_quiz) == 10:
                    st.session_state.quiz_data = new_quiz
                    st.session_state.user_answers = {}
                    st.session_state.current_index = 0
                    st.session_state.ai_hint_cache = {}
                    st.session_state.ai_solution_cache = {}
                    
                    # Sinkronisasi Sesi Baru ke Database Real-time
                    st.session_state.session_id = str(uuid.uuid4())
                    update_progress_siswa(
                        st.session_state.session_id, st.session_state.nama_siswa,
                        st.session_state.jenjang, st.session_state.mapel, 1, [], "BERJALAN"
                    )
                    
                    st.session_state.page = "quiz"
                    st.rerun()

    with col_act2:
        if st.button("⚙️ Pilih Mata Pelajaran Lain", use_container_width=True):
            st.session_state.page = "select_mapel"
            st.rerun()

    st.write("---")
    st.markdown("### 📖 Pembahasan Rinci dari Pembina RoboMANTAP ")
    st.caption("💡 *Untuk meminta RoboMANTAP membahas nya, Klik pada masing-masing soal di bawah ini ya!* 😊")
    
    for idx, q in enumerate(quiz_data):
        u_ans = user_answers.get(idx, "Tidak Dijawab")
        is_correct = u_ans == q["correct_answer"]
        status_icon = "✅ BENAR" if is_correct else ("❌ SALAH" if u_ans != "Tidak Dijawab" else "⚪ KOSONG")
        
        with st.expander(f"Soal No. {idx + 1} [{status_icon}] - Jawaban Anda: {u_ans}"):
            st.markdown(f"**Soal:**\n{q['question']}")
            st.markdown(f"**Kunci Jawaban:** {q['correct_answer']}")
            st.write("---")
            
            # Cache pembahasan per soal.
            solution_key = (
                st.session_state.mapel,
                idx,
                q["question"],
                q["correct_answer"],
            )

            if solution_key in st.session_state.ai_solution_cache:
                st.markdown("**🧕🏼 Pembahasan dari RoboMANTAP:**")
                st.markdown(st.session_state.ai_solution_cache[solution_key])

            if st.button(f"Tampilkan Pembahasannya dong! (Soal {idx + 1})", key=f"btn_sol_{idx}"):
                if solution_key in st.session_state.ai_solution_cache:
                    st.markdown("**🧕🏼 Pembahasan dari RoboMANTAP:**")
                    st.markdown(st.session_state.ai_solution_cache[solution_key])
                else:
                    st.markdown("**🧕🏼 Pembahasan dari RoboMANTAP:**")
                    streamed_solution = st.write_stream(
                        get_ai_solution_stream(
                            q["question"],
                            q["correct_answer"],
                            st.session_state.mapel,
                        ),
                        cursor="▌",
                    )

                    if streamed_solution and "⚠️" not in str(streamed_solution):
                        st.session_state.ai_solution_cache[solution_key] = str(streamed_solution)
